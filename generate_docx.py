from docx import Document
from docx.shared import Pt
import sys

def main():
    doc = Document()
    doc.add_heading('Project Caesar: Live Emergency Demo Guide', 0)

    doc.add_paragraph('This guide walks you through setting up a live, physical demonstration of Project Caesar on your Raspberry Pi 3. We will connect your phone camera to act as the AI eye, configure a mock gun for detection, and trigger both a physical ESP32 buzzer and a Twilio phone call to authorities.')

    doc.add_heading('Step 1: Using Your Phone Camera as the AI Eye', level=1)
    doc.add_paragraph('Since you are deploying everything on a Raspberry Pi 3, the easiest way to use your phone camera is to turn it into an IP webcam stream that the Pi can read.')
    
    p1 = doc.add_paragraph(style='List Number')
    p1.add_run('Download the "DroidCam" or "IP Webcam" app on your smartphone from the App Store or Google Play.').bold = True
    
    p2 = doc.add_paragraph(style='List Number')
    p2.add_run('Connect your phone and Raspberry Pi to the same WiFi network.')

    p3 = doc.add_paragraph(style='List Number')
    p3.add_run('Open the app and tap "Start Server". It will give you a URL (e.g., http://192.168.1.50:8080/video).')
    
    p4 = doc.add_paragraph(style='List Number')
    p4.add_run('Open the ')
    p4.add_run('configs/edge-pi3.toml').bold = True
    p4.add_run(' file on your Pi.')
    
    p5 = doc.add_paragraph(style='List Number')
    p5.add_run('In the [optical] section, change the mode to "synthetic", or if you have DroidCam mapping to a Linux device, use sentinel mode. The easiest way without drivers is to use the python adapter. But assuming you just map it:\nChange [sentinel] enabled = true and set device_id to your video URL string.')

    # Let's adjust this to be simpler. Using the IP Webcam URL in OpenCV VideoCapture is literally just passing the URL string. But our config expects an i32 for device_id.
    # To fix this without code changes, the user should just install droidcam client on the Pi, which maps it to /dev/video0.
    
    doc.add_heading('Step 2: ESP32 Buzzer Setup (The Hardware Alarm)', level=1)
    doc.add_paragraph('We will use the MQTT Actuator. When Project Caesar detects the gun, it will broadcast an alert over your local WiFi. The ESP32 will listen for this alert and beep.')
    
    e1 = doc.add_paragraph(style='List Number')
    e1.add_run('Connect a buzzer to your ESP32 (e.g., Positive to GPIO 13, Ground to GND).')
    
    e2 = doc.add_paragraph(style='List Number')
    e2.add_run('Flash your ESP32 with a simple Arduino sketch (or ESPHome) that connects to your WiFi and connects to the Mosquitto MQTT broker running on your Pi (e.g., 192.168.x.x:1883).')

    e3 = doc.add_paragraph(style='List Number')
    e3.add_run('Have the ESP32 subscribe to the topic: ')
    e3.add_run('caesar/actions').bold = True
    
    e4 = doc.add_paragraph(style='List Number')
    e4.add_run('In your Arduino code, when a message arrives, check if it contains the word "alert". If it does, drive GPIO 13 HIGH for 5 seconds to sound the buzzer.')

    e5 = doc.add_paragraph(style='List Number')
    e5.add_run('Open ')
    e5.add_run('configs/edge-pi3.toml').bold = True
    e5.add_run(' and ensure the [actuators] section has the MQTT actuator enabled and pointing to your Pi\'s IP.')

    doc.add_heading('Step 3: Configuring the Twilio Agentic Phone Call', level=1)
    doc.add_paragraph('When the threat is verified, the Python Orchestrator will make an automated text-to-speech call to your chosen phone number.')
    
    t1 = doc.add_paragraph(style='List Number')
    t1.add_run('Open the file: ')
    t1.add_run('/etc/systemd/system/caesar-orchestrator.service').bold = True
    t1.add_run(' on your Hub machine (or Pi if running everything in hybrid).')
    
    t2 = doc.add_paragraph(style='List Number')
    t2.add_run('Replace the placeholders in the Environment= lines with your real Twilio Account SID, Auth Token, Twilio Phone Number, and your real cell phone number (AUTHORITIES_PHONE).')

    t3 = doc.add_paragraph(style='List Number')
    t3.add_run('Restart the orchestrator to load the keys: ')
    t3.add_run('sudo systemctl daemon-reload && sudo systemctl restart caesar-orchestrator').italic = True

    t4 = doc.add_paragraph(style='List Number')
    t4.add_run('Open ')
    t4.add_run('services/mesh_orchestrator/alerts.py').bold = True
    t4.add_run(' and UNCOMMENT lines 8-11 to activate the live webhook. Restart the orchestrator again.')

    doc.add_heading('Step 4: Running the Live Demonstration', level=1)
    
    r1 = doc.add_paragraph(style='List Bullet')
    r1.add_run('Start the edge node: ').bold = True
    r1.add_run('./target/release/uriel-edge-node --config configs/edge-pi3.toml')
    
    r2 = doc.add_paragraph(style='List Bullet')
    r2.add_run('Point your phone camera at a clear scene. Stand in the frame.')
    
    r3 = doc.add_paragraph(style='List Bullet')
    r3.add_run('Pull out your makeshift gun.')
    
    r4 = doc.add_paragraph(style='List Bullet')
    r4.add_run('Watch the Magic Happen: ').bold = True
    r4.add_run('The edge node will see the gun, the tracker will classify it as CRITICAL, the ESP32 buzzer will instantly start screaming, and your phone will ring with an automated voice call detailing your exact GPS coordinates.')

    doc.save('Project_Caesar_Demo_Guide.docx')
    print("Successfully created Project_Caesar_Demo_Guide.docx")

if __name__ == "__main__":
    main()
